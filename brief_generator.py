#!/usr/bin/env python3
"""
Pre-Call Research Brief Generator
Powered by Claude AI — Nooks GTM Toolkit

SETUP (one time):
  pip install -r requirements.txt
  export ANTHROPIC_API_KEY=your_key_here   ← get it at console.anthropic.com

USAGE:
  python brief_generator.py
"""

import os
import sys
import re
import anthropic
import requests
from bs4 import BeautifulSoup
from datetime import datetime
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

# Optional: richer news context. Install with: pip install duckduckgo-search
try:
    from duckduckgo_search import DDGS
    HAS_SEARCH = True
except ImportError:
    HAS_SEARCH = False


def fetch_website(url):
    """Pull visible text from a company's homepage."""
    if not url:
        return ""
    try:
        if not url.startswith("http"):
            url = "https://" + url
        headers = {"User-Agent": "Mozilla/5.0 (Macintosh; Intel Mac OS X 10_15_7) AppleWebKit/537.36"}
        r = requests.get(url, headers=headers, timeout=10)
        soup = BeautifulSoup(r.text, "html.parser")
        for tag in soup(["script", "style", "nav", "footer", "header"]):
            tag.decompose()
        text = soup.get_text(separator=" ", strip=True)
        return text[:4000]
    except Exception as e:
        return f"(Could not fetch website: {e})"


def fetch_news(company_name):
    """Pull recent news headlines about the company (requires duckduckgo-search)."""
    if not HAS_SEARCH:
        return ""
    try:
        results = []
        with DDGS() as ddgs:
            for r in ddgs.news(company_name, max_results=5):
                results.append(f"- {r['title']} ({r.get('date', 'recent')}): {r['body'][:200]}")
        return "\n".join(results) if results else ""
    except Exception:
        return ""


def _google_search(query, label, num_results=10):
    """
    Search Google and return titles, URLs, and snippets.
    Uses requests + BeautifulSoup — no extra library needed.
    Handles Google's HTML with two fallback selector strategies.
    """
    from urllib.parse import quote as url_quote, unquote

    try:
        search_url = (
            f"https://www.google.com/search"
            f"?q={url_quote(query)}&num={num_results}&hl=en&gl=us"
        )
        headers = {
            "User-Agent": (
                "Mozilla/5.0 (Macintosh; Intel Mac OS X 10_15_7) "
                "AppleWebKit/537.36 (KHTML, like Gecko) "
                "Chrome/120.0.0.0 Safari/537.36"
            ),
            "Accept-Language": "en-US,en;q=0.9",
            "Accept": "text/html,application/xhtml+xml,application/xml;q=0.9,*/*;q=0.8",
        }

        r = requests.get(search_url, headers=headers, timeout=12)
        if r.status_code == 429:
            return f"[{label} — Google rate-limited, try again in a moment]"
        if r.status_code != 200:
            return f"[{label} — Google returned HTTP {r.status_code}]"

        soup = BeautifulSoup(r.text, "html.parser")
        results = []

        # Primary strategy: div.g result containers
        for container in soup.select("div.g"):
            h3 = container.find("h3")
            if not h3:
                continue
            title = h3.get_text(strip=True)

            a = container.find("a", href=True)
            href = a["href"] if a else ""
            if href.startswith("/url?q="):
                href = unquote(href.split("/url?q=")[1].split("&")[0])

            # Snippet — try several selector patterns Google has used
            snippet = ""
            for sel in ["div.VwiC3b", "div[data-sncf]", "span.aCOpRe",
                         "div.s", "div.IsZvec"]:
                el = container.select_one(sel)
                if el:
                    snippet = el.get_text(separator=" ", strip=True)[:220]
                    break

            if title:
                results.append(f"  • {title}\n    {href}\n    {snippet}")

        if results:
            return f"[{label} — {len(results)} result(s)]\n" + "\n".join(results)

        # Fallback: if Google's HTML changed, grab all h3 headings as titles
        fallback = [
            f"  • {h3.get_text(strip=True)}"
            for h3 in soup.find_all("h3")[:num_results]
            if len(h3.get_text(strip=True)) > 5
        ]
        if fallback:
            return (
                f"[{label} — {len(fallback)} result(s) (title-only fallback)]\n"
                + "\n".join(fallback)
            )

        return f"[{label} — 0 results]"

    except Exception as e:
        return f"[{label} — search error: {e}]"


def _fetch_page(url, label, char_limit=3000):
    """
    Directly fetch a URL and return visible text. Skips pages that redirect
    to a login wall (LinkedIn, Glassdoor) by checking for auth keywords.
    """
    try:
        headers = {
            "User-Agent": "Mozilla/5.0 (Macintosh; Intel Mac OS X 10_15_7) "
                          "AppleWebKit/537.36 (KHTML, like Gecko) Chrome/120.0.0.0 Safari/537.36"
        }
        r = requests.get(url, headers=headers, timeout=10, allow_redirects=True)
        if r.status_code != 200:
            return f"[{label} — HTTP {r.status_code}, skipped]"

        soup = BeautifulSoup(r.text, "html.parser")
        for tag in soup(["script", "style", "nav", "footer", "header"]):
            tag.decompose()
        text = soup.get_text(separator=" ", strip=True)

        # Detect login walls — return nothing useful rather than mislead Claude
        login_signals = ["sign in to view", "join to see", "log in to continue",
                         "create a free account", "authwall", "join linkedin"]
        if any(sig in text.lower() for sig in login_signals):
            return f"[{label} — login wall, skipped]"

        text = text[:char_limit]
        return f"[{label} — fetched OK]\n{text}" if text else f"[{label} — empty page]"
    except Exception as e:
        return f"[{label} — fetch error: {e}]"


def fetch_sdr_signals(company_name, website_url=""):
    """
    Surface SDR/BDR team signals via Google search + direct job-board page fetches.

    Layer 1 — Google searches (requests + BeautifulSoup, no extra library):
      Each query targets one concept. Google surfaces LinkedIn profiles, Glassdoor
      reviews, job postings, and news articles that DuckDuckGo misses.

    Layer 2 — Direct page fetches of public, login-free sources:
      Company careers page, Greenhouse board, Lever board.
      (The Org excluded — data can be outdated.)
    """
    slug = company_name.lower().replace(" ", "-").replace(".", "").replace(",", "").replace("'", "")
    findings = []

    # ── Layer 1: Google searches ──────────────────────────────────────────────

    # Core SDR/BDR title searches — catch LinkedIn profiles, Glassdoor, news
    findings.append(_google_search(
        f'{company_name} "sales development representative"',
        "Google 1 — SDR title"
    ))
    findings.append(_google_search(
        f'{company_name} "business development representative"',
        "Google 2 — BDR title"
    ))

    # LinkedIn specifically — Google indexes LinkedIn profile snippets even when DDG doesn't
    findings.append(_google_search(
        f'{company_name} "sales development representative" site:linkedin.com',
        "Google 3 — LinkedIn SDR profiles"
    ))
    findings.append(_google_search(
        f'{company_name} "sales development" site:linkedin.com',
        "Google 4 — LinkedIn sales development (broad)"
    ))

    # Glassdoor — employee reviews often mention SDR teams and headcounts
    findings.append(_google_search(
        f'{company_name} "sales development" site:glassdoor.com',
        "Google 5 — Glassdoor SDR mentions"
    ))

    # Job boards — Greenhouse and Lever indexed by Google
    findings.append(_google_search(
        f'{company_name} "sales development representative" site:boards.greenhouse.io',
        "Google 6 — Greenhouse SDR postings"
    ))
    findings.append(_google_search(
        f'{company_name} "sales development representative" site:jobs.lever.co',
        "Google 7 — Lever SDR postings"
    ))
    findings.append(_google_search(
        f'{company_name} "sales development representative" site:wellfound.com',
        "Google 8 — Wellfound SDR postings"
    ))

    # SDR leadership titles
    findings.append(_google_search(
        f'{company_name} "head of sales development" OR "VP of sales development" OR "director of sales development"',
        "Google 9 — SDR leadership (Head/VP/Director)"
    ))
    findings.append(_google_search(
        f'{company_name} "SDR manager" OR "BDR manager" OR "sales development manager"',
        "Google 10 — SDR manager titles"
    ))

    # Alternative outbound rep titles
    findings.append(_google_search(
        f'{company_name} "inside sales representative" OR "account development representative"',
        "Google 11 — ISR / ADR titles"
    ))

    # General hiring signal — catches any mention of SDR hiring on the open web
    findings.append(_google_search(
        f'{company_name} SDR hiring',
        "Google 12 — SDR hiring signal"
    ))
    findings.append(_google_search(
        f'{company_name} BDR hiring',
        "Google 13 — BDR hiring signal"
    ))

    # ── Layer 2: Direct page fetches (public, no login required) ─────────────

    # Company's own careers/jobs page — most authoritative for open SDR roles
    if website_url:
        domain = website_url.replace("https://", "").replace("http://", "").rstrip("/").split("/")[0]
        for path in ["/careers", "/jobs", "/about/careers", "/careers/open-roles"]:
            findings.append(_fetch_page(f"https://{domain}{path}", f"Careers page ({path})"))

    # Public Greenhouse job board — no login required
    findings.append(_fetch_page(
        f"https://boards.greenhouse.io/{slug}",
        "Greenhouse board (direct fetch)"
    ))

    # Public Lever job board — no login required
    findings.append(_fetch_page(
        f"https://jobs.lever.co/{slug}",
        "Lever board (direct fetch)"
    ))

    return "\n\n".join(f for f in findings if f) or "No SDR/BDR signals found across all sources."


def generate_brief(company_name, website_url, contact_name, contact_title):
    """Call Claude to generate the pre-call research brief."""
    website_content = fetch_website(website_url)
    news_content    = fetch_news(company_name)

    context_parts = []
    if website_content:
        context_parts.append(f"Website content:\n{website_content}")
    if news_content:
        context_parts.append(f"Recent news:\n{news_content}")
    context = "\n\n".join(context_parts) if context_parts else "No additional context available."

    prompt = f"""You are helping an Account Executive at Nooks prepare for a discovery call.

Nooks is an AI-powered sales assistant platform that helps B2B sales teams generate more pipeline from
outbound. It combines an AI parallel dialer, AI coaching assistant, and AI prospecting assistant into one
platform — helping SDR and BDR teams have more quality conversations, ramp faster, and hit quota.
Nooks customers typically see 2–3x pipeline per rep within days of adopting the platform.

Nooks sells to companies with dedicated SDR/BDR teams running high-volume outbound. The primary buyers
are VP of Sales, Head of Sales Development, CRO, and Revenue Operations leaders at B2B SaaS companies.

---
PROSPECT INFO:
Company:       {company_name}
Contact:       {contact_name or "Unknown"} — {contact_title or "Unknown title"}
---
CONTEXT:
{context}

---
Generate a tight, scannable pre-call research brief. Format it exactly like this:

## {company_name} — Pre-Call Brief

### What They Do
[2–3 sentences on the company's core business, go-to-market motion, and scale]

### Sales Team Signals
[Based on available signals, what does their outbound sales motion look like? Team size, SDR/BDR presence, tools they likely use (Outreach, Salesloft, Gong, ZoomInfo, etc.), hiring signals]

### Outbound Pain Points
[What pipeline generation and outbound challenges might a company like this face? Think: connect rates, rep ramp time, coaching at scale, tool sprawl, pipeline predictability]

### Nooks Fit
[How could Nooks specifically solve their problems — be concrete about the dialer, coaching, or prospecting angle]

### Tailored Discovery Questions
1. [Question]
2. [Question]
3. [Question]
4. [Question]
5. [Question]

### One-Line Opener
[A compelling, personalized opening line for the first 30 seconds of the call — sounds human, not scripted]

Keep the whole brief under 450 words. Make it something the AE can read in 90 seconds right before the call."""

    client = anthropic.Anthropic()
    message = client.messages.create(
        model="claude-opus-4-6",
        max_tokens=1200,
        messages=[{"role": "user", "content": prompt}]
    )
    return message.content[0].text


def generate_cold_emails(company_name, website_url, contact_name, contact_title, custom_notes):
    """Generate 3 cold email variations for a Nooks prospect."""
    website_content = fetch_website(website_url)
    news_content    = fetch_news(company_name)

    context_parts = []
    if website_content:
        context_parts.append(f"Website content:\n{website_content}")
    if news_content:
        context_parts.append(f"Recent news:\n{news_content}")
    if custom_notes:
        context_parts.append(f"Additional context from the rep:\n{custom_notes}")
    context = "\n\n".join(context_parts) if context_parts else "No additional context available."

    prompt = f"""You are helping an Account Executive at Nooks write cold outbound emails.

Nooks is an AI sales assistant platform that helps SDR and BDR teams generate more pipeline. It includes
an AI parallel dialer (2–3x more conversations per rep), an AI coaching assistant (faster ramp, better
conversion), and an AI prospecting assistant (less research time, higher intent outreach). Companies
switch to Nooks from fragmented stacks of dialers, sequencers, and coaching tools.

---
PROSPECT INFO:
Company:       {company_name}
Contact:       {contact_name or "Unknown"} — {contact_title or "Unknown title"}
---
CONTEXT:
{context}

---
Write 3 cold email variations. Each needs a subject line and a short body.

Format exactly like this:

---
**Variation 1: Direct**
**Subject:** [subject line]

[Body]

---
**Variation 2: Insight-Led**
**Subject:** [subject line]

[Body]

---
**Variation 3: Question-Based**
**Subject:** [subject line]

[Body]

---
Rules — follow every one of these:
- Each email must be under 100 words
- Use short paragraphs of 1 to 2 lines each
- Each email follows one simple structure: one observation about their sales team, one specific pain (low connect rates, rep ramp time, coaching at scale, pipeline unpredictability, or tool sprawl), one outcome Nooks delivers, one ask
- Never stack multiple features or products in the same email
- Use concrete, specific language. No "purpose-built", "robust", "seamlessly", or generic sales phrases
- Write like a peer who has seen this problem before at their scale, not like a salesperson pitching software
- No dashes used as punctuation. No overly polished transitions
- Sign off with: [Your name] at Nooks
- One CTA per email, kept light: "curious if you've run into this?" or "worth a quick call?" — never "book a demo"
- The tone should feel like it was written in 30 seconds, not edited for 30 minutes
"""

    client = anthropic.Anthropic()
    message = client.messages.create(
        model="claude-opus-4-6",
        max_tokens=1500,
        messages=[{"role": "user", "content": prompt}]
    )
    return message.content[0].text


def generate_crm_summary(call_notes, company_name, contact_name, contact_title, call_date):
    """Generate a CRM-ready call summary with standard fields and MEDDPICC scoring."""

    prompt = f"""You are helping an Account Executive at Nooks log a sales call in their CRM.

Nooks is an AI-powered sales assistant platform — AI parallel dialer, AI coaching assistant, and AI
prospecting assistant — for SDR and BDR teams at B2B SaaS companies.

---
CALL INFO:
Company:      {company_name or "Unknown"}
Contact:      {contact_name or "Unknown"} — {contact_title or "Unknown title"}
Date:         {call_date or "Unknown"}
---
CALL NOTES:
{call_notes}

---
Generate a structured CRM summary in two parts.

PART 1 — CALL SUMMARY (for the CRM notes field):

### Overview
[2–3 sentences capturing what was discussed and the outcome of the call]

### Pain Points
[Bullet list of specific pains the prospect mentioned — connect rates, ramp time, coaching gaps, tool sprawl, pipeline shortfalls, etc.]

### Objections
[Bullet list of objections raised, or "None raised" if none]

### Stakeholders Mentioned
[Name — Title — Role in deal, one per line. "None mentioned" if none]

### Next Steps
[Numbered list of agreed next steps with owners]

### Deal Stage
[Recommended CRM stage and one sentence of reasoning]

---

PART 2 — MEDDPICC SCORECARD:

**Metrics:** [What quantifiable outcomes did they mention? Pipeline per rep, connect rates, ramp time, quota attainment, number of conversations per day?]
**Economic Buyer:** [Who controls the budget? VP Sales, CRO, Head of SDR? Identified or unknown?]
**Decision Criteria:** [What will they use to evaluate solutions? Ease of use, dialer quality, AI accuracy, integration with their CRM/sequencer?]
**Decision Process:** [How will they decide? Timeline? Who else is involved? Is there a formal evaluation?]
**Paper Process:** [Any legal, procurement, or security review steps mentioned?]
**Identify Pain:** [What is the core business pain driving this evaluation? Low pipeline, poor rep productivity, high ramp time?]
**Champion:** [Is there an internal advocate pushing for this? Who?]
**Competition:** [Are they evaluating other tools? Outreach, Salesloft, Orum, Koncert, or other dialers mentioned?]

---
Rules:
- Be specific, not generic. Use exact words or numbers from the call notes where possible.
- If something wasn't mentioned, write "Not discussed" rather than guessing.
- Keep the tone professional and factual — this is going into a CRM, not a pitch deck.
- Total length should be under 450 words.
"""

    client = anthropic.Anthropic()
    message = client.messages.create(
        model="claude-opus-4-6",
        max_tokens=1500,
        messages=[{"role": "user", "content": prompt}]
    )
    return message.content[0].text


def generate_icp_score(company_name, website_url):
    """Research a company and score it against Nooks' ICP (1-4)."""
    website_content   = fetch_website(website_url)
    news_content      = fetch_news(company_name)
    linkedin_signals  = fetch_sdr_signals(company_name, website_url)

    context_parts = []
    if website_content:
        context_parts.append(f"WEBSITE CONTENT:\n{website_content}")
    if news_content:
        context_parts.append(f"RECENT NEWS:\n{news_content}")
    if linkedin_signals:
        context_parts.append(f"LINKEDIN SDR/BDR SIGNALS (searched LinkedIn via DuckDuckGo):\n{linkedin_signals}")
    context = "\n\n".join(context_parts) if context_parts else "No additional context available."

    prompt = f"""You are a sales researcher helping Nooks' GTM team decide whether to pursue a prospect.

Nooks' Ideal Customer Profile (ICP):
- B2B companies with a dedicated SDR or BDR team running high-volume outbound
- Industries: B2B SaaS, fintech, cybersecurity, HR tech, MarTech, sales tech, healthcare tech, logistics tech
- Company stage: Series B and beyond, or established companies scaling their outbound sales motion
- Team signals: 5+ SDRs/BDRs, active hiring for SDR/BDR roles, SDR leader (Head of Sales Dev / VP Sales Dev) present
- Pain signals: low connect rates on outbound calls, long rep ramp times (90+ days), inconsistent coaching, relying on manual research, scattered tool stack (separate dialer + sequencer + coaching tool)
- Decision makers: VP of Sales, Head of Sales Development, CRO, Revenue Operations

NOT a good fit:
- Companies with no outbound motion (inbound-only or PLG-only)
- Very early-stage startups with no dedicated SDR team yet
- Companies without a phone-heavy sales motion (purely email or social)
- Non-B2B companies (B2C, consumer, retail)
- Very small teams (fewer than 3 SDRs) with no growth plans

---
COMPANY: {company_name}
WEBSITE: {website_url or "Not provided"}
---
CONTEXT:
{context}

---
HOW TO INTERPRET THE SDR/BDR SIGNALS:
Up to 13 searches and 4 direct page fetches were run across job boards (Greenhouse, Lever, Wellfound),
Glassdoor, The Org, and the company's own careers page. Use all results together.

CRITICAL RULES — read carefully before scoring:

1. ZERO SEARCH RESULTS ≠ NO SDR TEAM.
   Many real SDR teams simply don't appear in web search results. Small or mid-market B2B SaaS
   companies often have 5–20 SDRs with zero public web footprint. If searches return nothing,
   mark SDR/BDR Team Size as "Unknown" — NOT "Weak" or "Poor fit."
   Base the score on industry fit, company stage, and go-to-market model instead.

2. ANY positive signal is strong evidence. One Glassdoor mention, one job posting, one Greenhouse
   listing, or one The Org result = confirmed SDR function. Weight it heavily.

3. Reason from company characteristics when web signals are thin:
   - B2B SaaS, Series B+, 50–500 employees, selling to enterprise or mid-market → almost certainly has SDRs
   - Cloud cost, security, HR tech, revenue ops, sales tech → SDR-heavy sectors
   - PLG/freemium, consumer, very early-stage → less likely to have SDRs
   Use this reasoning to fill gaps when search results are empty.

4. Alternative titles (ADR, MDR, ISR, inside sales rep, outbound sales rep) = SDR equivalent.
   Count them the same.

5. Leadership found (VP/Director/Head/Manager of Sales Dev) = strong signal of structured SDR org.

---
Research this company and produce an ICP scorecard. Format it exactly like this:

### ICP Score: [1, 2, 3, or 4] / 4

**[One-line verdict — e.g. "Strong fit — pursue now" or "Poor fit — deprioritize"]**

---

### Why This Score

[2–3 sentences explaining the reasoning. Be specific about signals found — especially what the LinkedIn data showed about their SDR/BDR team.]

---

### SDR / BDR Team Intelligence

- **Estimated SDR/BDR headcount:** [Best estimate from all search results — e.g. "~15 reps found across LinkedIn + Glassdoor" or "No individual reps found"]
- **Title variants found:** [List any non-standard titles surfaced — ADR, MDR, ISR, inside sales, etc.]
- **Leadership present:** [Yes / No / Unknown — name specific leaders found and their titles]
- **Active hiring:** [Yes / No — cite any specific job postings found]
- **Signal confidence:** [High / Medium / Low — based on how many of the 5 searches returned results]

---

### ICP Signal Breakdown

- **Outbound Sales Motion:** [Strong / Moderate / Weak / Unknown] — [one line explanation]
- **SDR / BDR Team Size:** [Strong / Moderate / Weak / Unknown] — [one line with the actual LinkedIn-derived estimate]
- **Industry Fit:** [Strong / Moderate / Weak / Unknown] — [one line explanation]
- **Company Stage & Growth:** [Strong / Moderate / Weak / Unknown] — [one line explanation]
- **Pain Signal Presence:** [Strong / Moderate / Weak / Unknown] — [one line explanation]

---

### Recommended Action

[One of these four, based on the score:]
- **Score 4:** Prioritize immediately. Add to active pipeline and reach out this week.
- **Score 3:** Worth pursuing. Research further and add to outbound sequence.
- **Score 2:** Possible fit. Monitor and revisit when you have more information.
- **Score 1:** Deprioritize. Move on — better opportunities exist.

---

Scoring guide:
- 4 = 4–5 strong ICP signals — clear fit, high priority
- 3 = 3 strong signals or 4–5 moderate ones — good fit, worth pursuing
- 2 = 1–2 strong signals or mixed signals — possible fit, needs qualification
- 1 = Few or no ICP signals — poor fit, not worth time right now

Be honest and precise. The LinkedIn data is ground truth for SDR team size — use it.
A bad lead wastes more time than no lead.
"""

    client = anthropic.Anthropic()
    message = client.messages.create(
        model="claude-opus-4-6",
        max_tokens=1200,
        messages=[{"role": "user", "content": prompt}]
    )
    return message.content[0].text


def parse_email_variations(emails_text):
    """Parse AI email output into a list of (title, subject, body) tuples."""
    variations = []
    chunks = re.split(r'\n?---\n?', emails_text)

    for chunk in chunks:
        chunk = chunk.strip()
        if not chunk:
            continue

        lines = chunk.split('\n')
        title, subject, body_lines, found_subject = "", "", [], False

        for line in lines:
            clean = line.strip()
            if re.match(r'\*\*Variation \d+', clean):
                title = re.sub(r'\*\*', '', clean).strip()
            elif clean.startswith('**Subject:**'):
                subject = clean.replace('**Subject:**', '').strip()
                found_subject = True
            elif found_subject:
                body_lines.append(line)

        if title and subject:
            body = '\n'.join(body_lines).strip()
            body = re.sub(r'\n{3,}', '\n\n', body)
            variations.append((title, subject, body))

    return variations


def save_as_docx(brief_text, filepath):
    """Convert the markdown-style brief into a formatted Word document."""
    doc = Document()

    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    for line in brief_text.splitlines():
        line = line.strip()

        if line.startswith("## "):
            p = doc.add_heading(line[3:], level=1)
            p.runs[0].font.color.rgb = RGBColor(0x7C, 0x3A, 0xED)  # Nooks purple

        elif line.startswith("### "):
            doc.add_heading(line[4:], level=2)

        elif re.match(r"^\*(.+)\*$", line):
            p = doc.add_paragraph()
            run = p.add_run(re.match(r"^\*(.+)\*$", line).group(1))
            run.italic = True
            run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)

        elif re.match(r"^\d+\.", line):
            p = doc.add_paragraph(style="List Bullet")
            p.style = doc.styles["Normal"]
            p.add_run(line)

        elif line == "" or line == "---":
            doc.add_paragraph("")

        else:
            doc.add_paragraph(line)

    doc.save(filepath)


def main():
    print("\n╔══════════════════════════════════════╗")
    print("║   Pre-Call Research Brief Generator  ║")
    print("╚══════════════════════════════════════╝\n")

    if not os.environ.get("ANTHROPIC_API_KEY"):
        print("⚠️  ANTHROPIC_API_KEY not set.")
        print("   Get your key at: https://console.anthropic.com")
        print("   Then run:  export ANTHROPIC_API_KEY=your_key_here\n")
        sys.exit(1)

    company_name  = input("Company name:   ").strip()
    if not company_name:
        print("Company name is required.")
        sys.exit(1)

    website_url   = input("Website URL:    ").strip()
    contact_name  = input("Contact name:   ").strip()
    contact_title = input("Contact title:  ").strip()

    print(f"\n⏳ Researching {company_name}...")
    if HAS_SEARCH:
        print("   → Fetching website & recent news...")
    else:
        print("   → Fetching website content...")
    print("   → Generating brief with Claude...\n")

    try:
        brief = generate_brief(company_name, website_url, contact_name, contact_title)
    except anthropic.AuthenticationError:
        print("❌ Invalid API key. Double-check your ANTHROPIC_API_KEY.")
        sys.exit(1)
    except Exception as e:
        print(f"❌ Unexpected error: {e}")
        sys.exit(1)

    print("─" * 50)
    print(brief)
    print("─" * 50)

    desktop     = os.path.join(os.path.expanduser("~"), "Desktop")
    safe_name   = company_name.replace("/", "").strip()
    folder_path = os.path.join(desktop, safe_name)
    os.makedirs(folder_path, exist_ok=True)

    filename = os.path.join(folder_path, "brief.docx")
    save_as_docx(brief, filename)

    print(f"\n✅ Saved to Desktop → {safe_name} → brief.docx")
    print("   Open it in Word or drag it into Google Docs.\n")

if __name__ == "__main__":
    main()
